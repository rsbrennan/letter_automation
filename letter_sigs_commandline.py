from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#import io
#document = Document()
import pandas as pd

import argparse


parser = argparse.ArgumentParser()
parser.add_argument("-i", "--input",
	help="specify the path to the input excel file here.",
	required=True) # specify which command-line options the program is willing to accep

parser.add_argument("-s", "--style", 
	type=argparse.FileType('r'),
	help="specify the file and path to the style docx file here.",
	required=True)

parser.add_argument("-t", "--title",
	help="title under each name",
	required=True)

parser.add_argument("-o", "--output",
	help="specify the file and path of the out word doc here. should end in .docx",
	required=True)

args = parser.parse_args()

# call any argument with args.i for example

# python -m pip install pandas
# python -m pip install xlrd

#### important things ###
# no spaces in any file names
# make sure that the names in the excel doc are in the sheet called Sheet1
# include a "Preferred_Name" column for the name you want listed in the docx file


form_path = args.input

xl = pd.ExcelFile(form_path)
#print(xl.sheet_names)

df = pd.read_excel(xl, 'Sheet1')

# get number of people on letter
ct_len = df.Preferred_Name.count()

# add col of index
df['indx_col'] = range(1, len(df) + 1)

# check that I actually want to add 1
even_indexes = range(2,ct_len+1,2)
odd_indexes = range(1,ct_len+1,2)

df_even = df[df['indx_col'].isin(even_indexes)]
df_odd = df[df['indx_col'].isin(odd_indexes)]

# add new sequential index

df_even = df_even.copy()
df_odd = df_odd.copy()
df_even['indx_col'] = range(1, len(df_even) + 1)
df_odd['indx_col'] = range(1, len(df_odd) + 1)



#### make word doc

# read in word doc that has correct table format
document = Document(args.style)

table = document.add_table((len(df)/2)*7, 3, style='CustomTable')

# first column:

# adding by row then col. for this,
row_start = 0
for index, row in df_odd.iterrows():
	#print row['Preferred_Name']
	#print index
	table.cell(row_start,0).text = '____________________'
	table.cell(row_start+1,0).text = row['Preferred_Name']
	table.cell(row_start+2,0).text = args.title
	row_start = row_start + 6

row_start = 0
for index, row in df_even.iterrows():
	#print row['Preferred_Name']
	#print index
	table.cell(row_start,2).text = '____________________'
	table.cell(row_start+1,2).text = row['Preferred_Name']
	table.cell(row_start+2,2).text = args.title
	row_start = row_start + 6

document.save(args.output)

